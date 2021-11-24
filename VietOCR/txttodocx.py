from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from docx.shared import Inches

inputsDir = 'texts/Tu dien Hoang Phe/results'
path = list(filter(lambda x: x.endswith("txt"), os.listdir(inputsDir)))
# x = path[3]
# print(path[0][5:7], x[8:9], x[10:10 + min(x[10:].find('-') if x[10:].find('-') != -1 else 2000, x[10:].find('.'))])
path = sorted(path, key = lambda x: int(x[:7])*10000 + int(x[8:9])*1000 + int(x[10:10 + min(x[10:].find('-') if x[10:].find('-') != -1 else 2000, x[10:].find('.'))]))
doc = Document()
para = ''
docPara = doc.add_paragraph('')

# Tags và Notes dùng cho từ điển 001 (Hoàng Phê):
tags = {' d.': ' danh từ', 'đg.': 'động từ', ' t.': ' tính từ', 'đ.': 'đại từ', ' p.': ' phụ từ', 'k.': 'kết từ', 'tr.': 'trợ từ', ' c.': ' cảm từ', 'hoài nghi đt.': 'hoài nghi động từ'}

notes = {'(id.).': '(ít dùng)', '(kng.).': '(khẩu ngữ)', '(ph.).': '(phương ngữ)', '(vch.).': '(văn chương)'}

# # Tags và Notes dùng cho từ điển 002 (Ng Kim Than):
# tags = { 'cd.': 'ca dao', 'dt.': 'danh từ', 'đt.': 'động từ', 'gt.': 'giới từ', 'id.': 'ít dùng', 'lt.': 'liên từ', ' ng.': ' nghĩa', 'pt.': 'phụ từ', 'tht.': 'thán từ', 'tng.': 'tục ngữ', 'trt.': 'trợ từ', 'vt.': 'vị từ'}

# notes = {'(id.).': '(ít dùng)', '(kng.)': '(khẩu ngữ)' , '(thgt.)': '(thông tục)', '(ph.)': '(phương ngữ)', '(vchg.)': '(văn chương)',  '(trtr.)': '(trang trọng)', '(kc.)': '(kiểu cách)', '(chm.)': '(chuyên môn)'} # còn nữa

# tags = {' dt.': ' danh từ', 'đgt.': 'động từ', ' tt.': ' tính từ', ' pht.': ' phụ từ',}

# notes = {'(id.).': '(ít dùng)', '(kng.).': '(khẩu ngữ)', '(ph.).': '(phương ngữ)', '(vch.).': '(văn chương)'}

alphabets = {'a': ['a', 'ă', 'â', 'à', 'á', 'ã', 'ả', 'ạ', 'ắ', 'ằ', 'ẵ', 'ẳ', 'ặ', 'ấ', 'ầ', 'ẫ', 'ẩ', 'ậ'], 'b': ['b'], 'c': ['c'], 'd': ['d', 'đ'], 'e': ['e', 'ê', 'é', 'è', 'ẽ', 'ẻ', 'ẹ', 'ế', 'ề', 'ễ', 'ể', 'ệ'], 'f': ['f'], 'g': ['g'], 'h': ['h'], 'i': ['i', 'í', 'ì', 'ĩ', 'ỉ', 'ị'], 'j': ['j'], 'k': ['k'], 'l': ['l'], 'm': ['m'], 'n': ['n'], 'o': ['o', 'ô', 'ơ', 'ó', 'ò', 'õ', 'ỏ', 'ọ', 'ố', 'ồ', 'ỗ', 'ổ', 'ộ', 'ớ', 'ờ', 'ỡ', 'ở', 'ợ'], 'p': ['p'], 'q': ['q'], 'r': ['r'], 's': ['s'], 't': ['t'], 'u': ['u', 'ư', 'ú', 'ù', 'ũ', 'ủ', 'ụ', 'ứ', 'ừ', 'ữ', 'ử', 'ự'], 'v': ['v'], 'w': ['w'], 'x': ['x'], 'y': ['y'], 'z': ['z'], '`':['`']}

extrasBeforeTag = {'cn.': '(cũng nói)', 'cv.': '(cũng viết)'}  # tiếp theo sẽ là 1 từ in nghiêng có chấm cuối câu 

extrasAfterTag = {' x.': ' Xem', 'Như': 'Như'}
newEntry = True
currentAlphabet = '`'
lastLine = '.'
pages = []
lastWord = ''
lastType = ''

def hasSubString(text, tags):
    for tag in tags.keys():
        if tag in text:
            return True
    return False

def hasExtraTag(text):
    for tag in extrasAfterTag.keys():
        if tag in text:
            index = text.find(tag)
            if index > 1 and text[index - 1] != '.' or text[index - 1] != '!' or text[index - 1] != '?' or text[index - 1] != ',':
                return True

    return False

def checkFirstWord(text, currAlphabet, alphabetsDict):
    for char in alphabetsDict[currAlphabet]:
        if text.lower().startswith(char) or text.lower().startswith('"' + char):
            return True
    return False

def checkOrderNumber(text):
    for ordNum in ['IV.', 'V.', 'III.', 'II.', 'I.']:
        idx = text.find(ordNum)
        if idx != - 1:
            return ordNum, idx
    return None

def checkNumber(text: str):
    for num in ['1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.']:
        idx = text.find(num)
        if idx != - 1:
            return num, idx
    return None

def formatCurrentWord(currWord: str):
    specialChars = ["?", "%", "("]
    for specialChar in specialChars:
        if currWord.startswith(specialChar):
            currWord = currWord[:0] + '"' + currWord[0+1:]
            currWord = currWord[:-1] + '"'
            break

    i = 0
    try:
        if currWord[i] not in [currentAlphabet, '"']:
            currWord = currWord[:i] + currentAlphabet + currWord[i+1:]
            # i += 1
    except Exception as _:
        pass
    return currWord

countNextAlphabet = 0
for txt in path:
    print(txt)
    txtFile = open(inputsDir + '/' + txt, encoding='utf-8', errors='ignore').read()
    txtFile = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+','', txtFile) # remove all non-XML-compatible characters
    txtFile = txtFile.replace('Ä', 'A')
    txtFile = txtFile.replace('ä', 'a')
    txtFile = txtFile.replace('“', '"')
    txtFile = txtFile.replace('”', '"')
    txtFile = txtFile.replace('¿', 't')
    txtFile = txtFile.replace(':.', 't.')
    txtFile = txtFile.replace('¡(', 't')
    
    for key, value in notes.items():
        txtFile = txtFile.replace(key, value)

    for key, value in extrasBeforeTag.items():
        txtFile = txtFile.replace(key, value)

    if int(txt[10:10 + min(txt[10:].find('-') if txt[10:].find('-') != -1 else 2000, txt[10:].find('.'))]) == 0 and int(txt[8:9]) == 0:
        docPara = doc.add_paragraph('')
        docPara.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docPara.add_run("TRANG " + txt[3:7]).bold = True

    line = txtFile.split('\n')[0]

    if newEntry or line.lower().startswith(chr(ord(currentAlphabet) + 1) + ',' + chr(ord(currentAlphabet) + 1)) or checkOrderNumber(line) or checkNumber(line):
        nextAlphabet = chr(ord(currentAlphabet) + 1)
        if line.lower().startswith(nextAlphabet + ',' + nextAlphabet):
            currentAlphabet = nextAlphabet
            docPara = doc.add_paragraph('')
            docPara.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            docPara.add_run(formatCurrentWord(nextAlphabet + ',' + nextAlphabet.upper() + ' ')).bold = True
            lastWord = nextAlphabet + ',' + nextAlphabet.upper() + ' '
            lastType = ''
            docPara.add_run(line[4:])

        else:
            if line.lower().startswith(nextAlphabet) and newEntry:
                countNextAlphabet += 1
                if countNextAlphabet == 3:
                    currentAlphabet = nextAlphabet
                    countNextAlphabet = 0
            else:
                countNextAlphabet = 0

            if (hasSubString(line.replace(',', '.').lower(), tags) or hasExtraTag(line.replace(',', '.'))): # mục từ mới)
                # if line.lower().startswith(chr(ord(currentAlphabet) + 1)) or line.lower().startswith('"' + chr(ord(currentAlphabet) + 1)):
                #     currentAlphabet = chr(ord(currentAlphabet) + 1)
                
                if (not checkFirstWord(line, currentAlphabet, alphabets)) and checkOrderNumber(line) == '':
                    if lastLine.endswith(' '):
                        docPara.add_run(line)
                    else:
                        docPara.add_run(' ' + line)
                    continue
                currentTags = {}
                
                tempLine = line.replace(',', '.').lower()
                for extra in extrasAfterTag.keys():
                    if extra in line.replace(',', '.'):
                        currentTags[extra] = line.find(extra)
                for tag in tags.keys():
                    if tag in tempLine:
                        currentTags[tag] = tempLine.find(tag)
                sortedCurrentTags = sorted(currentTags.items(), key=lambda kv: kv[1])
                word = line[:sortedCurrentTags[0][1]]
                # content = line[sortedCurrentTags[len(sortedCurrentTags) - 1][1] + len(sortedCurrentTags[len(sortedCurrentTags) - 1][0]):]
                content = line[sortedCurrentTags[-1][1] + len(sortedCurrentTags[-1][0]):]

                docPara = doc.add_paragraph('')
                docPara.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                docPara.add_run(formatCurrentWord(word)).bold = True
                lastWord = word

                # if sortedCurrentTags[0][0] in tags.keys():
                #     docPara.add_run(tags[sortedCurrentTags[0][0]]).italic = True
                #     lastType = tags[sortedCurrentTags[0][0]]
                # else:
                #     docPara.add_run(extrasAfterTag[sortedCurrentTags[0][0]]).italic = True
                #     lastType = extrasAfterTag[sortedCurrentTags[0][0]]

                for tagKey,_ in sortedCurrentTags:
                    if tagKey in tags.keys():
                        if tagKey.startswith(' '):
                            docPara.add_run(tags[tagKey]).italic = True
                            lastType += tags[tagKey]
                        else:
                            docPara.add_run(' ' + tags[tagKey]).italic = True
                            lastType += ' ' + tags[tagKey]
                    else:
                        docPara.add_run(' ' + extrasAfterTag[tagKey]).italic = True
                        lastType += ' ' + extrasAfterTag[tagKey]
                docPara.add_run(content)

            elif checkFirstWord(line, currentAlphabet, alphabets):
                docPara = doc.add_paragraph('')
                docPara.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                for i in range(2, len(line)):
                    # if (line[i].isupper() or line[i].isdigit()) and line[i - 1] == ' ' and (line[i - 2] != '.' and line[i - 2] != ',' and line[i - 2] != '!' and line[i - 2] != '?'):
                    if line[i].isupper() and line[i - 1] == ' ' and (line[i - 2] != '.' and line[i - 2] != ',' and line[i - 2] != '!' and line[i - 2] != '?'):
                        docPara.add_run(formatCurrentWord(line[:i])).bold = True
                        lastWord = line[:i]
                        docPara.add_run(line[i:])
                        lastType = ''
                        break
                else:
                    docPara.add_run(line)
            elif checkOrderNumber(line):
                num, idx = checkOrderNumber(line)
                docPara.add_run(line[:idx - 1])
                docPara = doc.add_paragraph('')
                docPara.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                docPara.add_run(formatCurrentWord(lastWord)).bold = True
                docPara.add_run(line[idx - 1:])
                lastType = ''
            
            elif checkNumber(line):
                num, idx = checkNumber(line)
                if num == '1.':
                    docPara.add_run(line[:line.find('.') + 1])
                    docPara = doc.add_paragraph('')
                    docPara.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    docPara.add_run(formatCurrentWord(line[line.find('.') + 1:idx - 1])).bold = True
                    docPara.add_run(line[idx - 1:])
                    lastWord = line[line.find('.') + 1:idx - 1]
                    lastType = ''
                else:
                    docPara.add_run(line[:idx - 1])
                    docPara = doc.add_paragraph('')
                    docPara.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    docPara.add_run(formatCurrentWord(lastWord)).bold = True
                    docPara.add_run(lastType).italic = True
            else:
                docPara = doc.add_paragraph('')
                docPara.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                docPara.add_run(line)
                lastWord = ''
                lastType = ''
        newEntry = False
    elif lastLine.endswith(' '):
        docPara.add_run(line)
    else:
        docPara.add_run(' ' + line)
    
    if txt[:-4].endswith('endEntry'):
        newEntry = True
        lastLine = line

doc.save('result.docx')
