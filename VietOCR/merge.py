import re
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
para = ''
docPara = doc.add_paragraph('')

contents = ''
i = 1
inputsDir = 'results'
path = os.listdir(inputsDir)
path = sorted(path, key = lambda x: len(x))

for txt in path:
    print(txt)
    content = open(inputsDir + '/' + txt, encoding='utf-8', errors='ignore').read()
    content = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+','', content) # remove all non-XML-compatible characters
    # contents += '[Slide ' +str(i) + ']\n'
    contents += (content)
    # docPara = doc.add_paragraph('')
    # docPara.add_run('[Slide ' +str(i) + ']').bold = True
    # docPara.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # docPara = doc.add_paragraph(content)
    # i += 1

# doc.save('result.docx')

with open('result.txt', 'w', encoding='utf-8') as f:
    f.write(contents)
